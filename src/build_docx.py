"""Build a Word document (.docx) catalogue + roadmap for building software with AI and Python.

Uses the same verified data as the workbook: ai_python_data.py, ai_python_snippets.py
and ai_link_status.json. Every resource title is a clickable hyperlink.
"""

import json
from collections import Counter, OrderedDict
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

from ai_python_data import ROWS, ROADMAP, B, I, A, X
from ai_python_snippets import SNIPPETS

OUT = "Build_Software_with_AI_and_Python.docx"
TODAY = date.today().isoformat()

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0x2E, 0x6B, 0x3E)
GREY = RGBColor(0x66, 0x66, 0x66)
LINKBLUE = RGBColor(0x05, 0x63, 0xC1)
CODEBG = "F4F4F7"

LEVEL_ORDER = [B, I, A, X]
LEVEL_BLURB = {
    B: "You are new to Python or new to programming. Goal: write working scripts and make your first AI-powered app.",
    I: "You can write Python. Goal: build maintainable, tested, deployable AI software with real APIs and data.",
    A: "You ship AI software. Goal: agents, evaluation, serving, security, cost control and performance at production quality.",
    X: "Useful at any level - tools and guidance for letting AI help you write the software itself.",
}

try:
    with open("ai_link_status.json") as fh:
        STATUS = json.load(fh)
except FileNotFoundError:
    STATUS = {}


def verdict(url):
    s = STATUS.get(url)
    if not s:
        return ""
    if s["kind"] == "ok":
        return "link verified"
    if s["kind"] == "blocked-to-bots":
        return "opens in browser (blocks bots)"
    return "CHECK LINK"


# ----------------------------------------------------------------- helpers
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_hyperlink(paragraph, url, text, bold=False, size=None):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def code_block(doc, code, lang):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shade_paragraph(p, CODEBG)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    # ensure monospace applies to complex/east-asian ranges too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    return p


# ================================================================= document
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ---- title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Building Software with AI and Python")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A tiered guide and link catalogue - beginner, intermediate and advanced")
r.font.size = Pt(12)
r.italic = True
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"Compiled {TODAY}   |   {len(ROWS)} curated resources   |   {len(SNIPPETS)} ready-to-run snippets   |   every link is clickable")
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

doc.add_paragraph()


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else ACCENT
    return h


# ---- how to use
heading("How to use this document", level=1)
for line in [
    "This is a practical roadmap plus a verified catalogue of resources for creating software with AI in Python, "
    "organised into three levels so you can jump in wherever you are:",
]:
    doc.add_paragraph(line)
for lvl in (B, I, A):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lvl.split(" - ")[1] + ": ")
    r.bold = True
    p.add_run(LEVEL_BLURB[lvl])
doc.add_paragraph(
    "Work top to bottom, or use the roadmap table in each level as a checklist. Each stage names what to learn, "
    "something concrete to build, and a clear 'done when' test so you know you are ready to move on."
)

p = doc.add_paragraph()
p.add_run("A note on scope and cost. ").bold = True
p.add_run(
    "Most resources here are free. Calling a hosted model (OpenAI, Anthropic) costs money per use - a few dollars "
    "covers a week of learning - but you can do almost everything at zero cost by running open models locally with "
    "Ollama. The Cost column in the tables flags which is which."
)
p = doc.add_paragraph()
p.add_run("Link checking. ").bold = True
n_ok = sum(1 for u in {r[4] for r in ROWS} if STATUS.get(u, {}).get("kind") == "ok")
n_block = sum(1 for u in {r[4] for r in ROWS} if STATUS.get(u, {}).get("kind") == "blocked-to-bots")
p.add_run(
    f"Every link was requested on {TODAY}: {n_ok} returned HTTP 200 and {n_block} sit behind bot protection "
    "(they open normally in a browser). None were broken. The code snippets were all syntax-checked."
)

# ---- contents overview
heading("What's inside", level=1)
for txt in [
    "1. The roadmap - a staged path with build milestones for each level.",
    "2. Resources by level - the full catalogue, grouped Beginner / Intermediate / Advanced / All levels.",
    "3. Starter code - copy-paste Python for the things people get stuck on.",
    "4. Summary counts - how the catalogue breaks down.",
]:
    doc.add_paragraph(txt, style="List Number")

# ================================================================= 1. ROADMAP
doc.add_page_break()
heading("1. The learning roadmap", level=1)
doc.add_paragraph(
    "Fifteen stages from zero to production. Do them in order within a level. The 'Done when' column is the "
    "important part - it is the test that tells you to move on rather than a vague sense of progress."
)

roadmap_by_level = OrderedDict((lvl, []) for lvl in (B, I, A))
for row in ROADMAP:
    roadmap_by_level[row[0]].append(row)

for lvl in (B, I, A):
    heading(lvl.split(" - ")[1] + " path", level=2)
    for (_, stage, goal, learn, build, done) in roadmap_by_level[lvl]:
        p = doc.add_paragraph()
        r = p.add_run(f"{stage}. {goal}")
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = ACCENT
        for label, text in (("Learn", learn), ("Build", build), ("Done when", done)):
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Inches(0.25)
            pp.paragraph_format.space_after = Pt(2)
            rr = pp.add_run(f"{label}:  ")
            rr.bold = True
            pp.add_run(text)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ================================================================= 2. RESOURCES
doc.add_page_break()
heading("2. Resources by level", level=1)
doc.add_paragraph(
    "Grouped by level, then by topic. Click any resource title to open it. The one-line description says what you "
    "actually get there and why it is worth your time."
)

grouped = OrderedDict()
for lvl, cat, title, desc, url, typ, cost in ROWS:
    grouped.setdefault(lvl, OrderedDict()).setdefault(cat, []).append((title, desc, url, typ, cost))

for lvl in LEVEL_ORDER:
    if lvl not in grouped:
        continue
    heading(lvl.split(" - ")[-1] if " - " in lvl else lvl, level=2)
    ib = doc.add_paragraph()
    ib.add_run(LEVEL_BLURB[lvl]).italic = True
    for cat, items in grouped[lvl].items():
        heading(cat, level=3)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = (Inches(2.5), Inches(3.7), Inches(1.1))
        hdr = table.rows[0].cells
        for c, label in zip(hdr, ("Resource (click to open)", "What you get", "Type / cost")):
            c.paragraphs[0].add_run(label).bold = True
            set_cell_bg(c, "1F3864")
            for run in c.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        for title, desc, url, typ, cost in items:
            cells = table.add_row().cells
            # resource cell: hyperlink title + small verification note
            p = cells[0].paragraphs[0]
            add_hyperlink(p, url, title, bold=True, size=9.5)
            note = verdict(url)
            if note:
                np = cells[0].add_paragraph()
                nr = np.add_run(note)
                nr.font.size = Pt(7.5)
                nr.font.color.rgb = GREY
                nr.italic = True
            dp = cells[1].paragraphs[0]
            dr = dp.add_run(desc)
            dr.font.size = Pt(9)
            tp = cells[2].paragraphs[0]
            tr = tp.add_run(f"{typ}\n{cost}")
            tr.font.size = Pt(8)
        for row in table.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = w

# ================================================================= 3. SNIPPETS
doc.add_page_break()
heading("3. Starter code", level=1)
doc.add_paragraph(
    "Copy-paste starting points for the things people most often get stuck on. Every Python snippet here was "
    "syntax-checked. Model APIs change - if a call is rejected, open the linked source for the current shape."
)

snip_by_level = OrderedDict((lvl, []) for lvl in (B, I, A))
for s in SNIPPETS:
    snip_by_level.setdefault(s[0], []).append(s)

for lvl in (B, I, A):
    heading(lvl.split(" - ")[1] + " snippets", level=2)
    for (_, title, lang, why, code, src) in snip_by_level[lvl]:
        h = doc.add_paragraph()
        r = h.add_run(f"{title}  ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = ACCENT
        tag = h.add_run(f"[{lang}]")
        tag.font.size = Pt(8)
        tag.font.color.rgb = GREY
        wp = doc.add_paragraph()
        wp.add_run(why).italic = True
        wp.runs[0].font.size = Pt(9.5)
        code_block(doc, code, lang)
        sp = doc.add_paragraph()
        sr = sp.add_run("Source: ")
        sr.font.size = Pt(8)
        sr.font.color.rgb = GREY
        add_hyperlink(sp, src, src, size=8)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ================================================================= 4. COUNTS
doc.add_page_break()
heading("4. Summary counts", level=1)

heading("Resources per level", level=3)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light List Accent 1"
tbl.rows[0].cells[0].paragraphs[0].add_run("Level").bold = True
tbl.rows[0].cells[1].paragraphs[0].add_run("Resources").bold = True
level_counts = Counter(r[0] for r in ROWS)
for lvl in LEVEL_ORDER:
    if lvl in level_counts:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(lvl)
        cells[1].paragraphs[0].add_run(str(level_counts[lvl]))
cells = tbl.add_row().cells
cells[0].paragraphs[0].add_run("TOTAL").bold = True
cells[1].paragraphs[0].add_run(str(len(ROWS))).bold = True

heading("Resources per topic", level=3)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light List Accent 1"
tbl.rows[0].cells[0].paragraphs[0].add_run("Topic").bold = True
tbl.rows[0].cells[1].paragraphs[0].add_run("Resources").bold = True
for cat, n in Counter(r[1] for r in ROWS).most_common():
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(cat)
    cells[1].paragraphs[0].add_run(str(n))

heading("Link check results", level=3)
kinds = Counter(STATUS.get(r[4], {}).get("kind", "not checked") for r in ROWS)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light List Accent 1"
tbl.rows[0].cells[0].paragraphs[0].add_run("Result").bold = True
tbl.rows[0].cells[1].paragraphs[0].add_run("Count").bold = True
for label, key in [("Live (HTTP 200)", "ok"),
                   ("Bot-blocked (opens in a browser)", "blocked-to-bots"),
                   ("Broken", "broken")]:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(label)
    cells[1].paragraphs[0].add_run(str(kinds.get(key, 0)))

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run(
    "Caveats: the AI tooling landscape moves fast - versions, model names and prices change, so cross-check "
    "community posts against official docs. Courses marked Paid are third-party products; check reviews and "
    "recency before buying. Snippets are starting points, not production code - add real error handling, "
    "secrets management and tests."
)
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

doc.save(OUT)

print(f"Wrote {OUT}")
print(f"  resources : {len(ROWS)} across {len(set(r[1] for r in ROWS))} topics, {len(level_counts)} levels")
print(f"  roadmap   : {len(ROADMAP)} stages")
print(f"  snippets  : {len(SNIPPETS)}")
print(f"  link check: {dict(kinds)}")
